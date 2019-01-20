from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import warnings
warnings.filterwarnings(action='ignore', category=UserWarning, module='gensim')
import gensim
from gensim.summarization.summarizer import summarize
from gensim.summarization import keywords
import requests
from win10toast import ToastNotifier
import os, PyPDF2
import time
from time import mktime, sleep
import textwrap3

toaster = ToastNotifier()

print("Synopsize [Version 1.0] | © 2019 | All rights reserved. ")
def main():
    print("\n")
    type_=input("What do you want to Synopsize?\n\n1. A bunch of text.\n2. A PDF Document.\n\nType: ")
    print("\n")
    if type_=="1" or type_=="A bunch of text".lower():
        type_1=input("Just one more thing, are you:\n\n1. Pasting a block of text or,\n2. Pasting a txt file?\n\nAnswer: ")
        if type_1=="1" or type_1=="Pasting a block of text".lower():
            print("Please ensure that you are pasting one block of text and not paragraphs. ")
            print("\n")
            text_string=input("Paste the block of text that requires a summary here: ")
            print("\n")
            ratio=input("How lengthy of a summary do you want?\n1. 20% of entire block\n2. 50% of entire block.\n\n1 or 2?  ")
            if ratio=="1" or ratio=="20%" or ratio =="20":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.2))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize20.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                main()
            if ratio=="2" or ratio=="50%" or ratio =="50":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.5))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize50.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                main()
        if type_1=="2" or type_1=="A txt file".lower() or type_1=="Block".lower() or type_1=="txt".lower() or type_1=="Txt file".lower():
            doc=input("Enter the document name in this format 'doc.txt': ")
            document_text = open(doc, 'r', encoding='utf-8')
            text_string = document_text.read().lower()
            ratio=input("How lengthy of a summary do you want?\n1. 20% of entire block\n2. 50% of entire block.\n\nSummary level: ?  ")
            if ratio=="1" or ratio=="20%" or ratio =="20":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.2))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize20.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                main()
            if ratio=="2" or ratio=="50%" or ratio =="50":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.5))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize500.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                main()
        else:
            if __name__ == "__main__":
                main()
            
    elif type_=="2" or type_=="A PDF Document".lower():
        doc=input("NOTE\n1. Please ensure Synopsize is in the same directory as the PDF.\n2. Enter the document name in this format 'doc.pdf'.\n3. Ensure document is readable (i.e in OCR format) and the document name has .pdf at the end otherwise I will abruptly close, and I'll leave you very angry :)** -  ")
        list=textwrap3.wrap(doc, width=95)
        for element in list:
            print("\n")
            print("About to summarize", element, "...")
        print("\n")
        def pdf_to_txt(file_path):
            #doc="cms.pdf"
            toaster.show_toast("Now converting PDF",
                                        "If your document is large, it may take some time.",
                                        icon_path="cap.ico",
                                        duration=5)
            fp = open(file_path, 'rb')
            rsrcmgr = PDFResourceManager()
            retstr = StringIO()
            codec = 'utf-8'
            laparams = LAParams()
            
            device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            for page in PDFPage.get_pages(fp, pagenos=set(), maxpages=0, password="", caching=True, check_extractable=True):
                interpreter.process_page(page)

            text = retstr.getvalue()

            fp.close()
            device.close()
            retstr.close()
            return text

        document_text= pdf_to_txt(doc)
        text_string = document_text
        toaster.show_toast("Almost done.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
        conversion=input("Stage one complete. Proceed to summary? [y/n] ")
        print("\n")
        if conversion=="y":
            ratio=input("How lengthy of a summary do you want?\n1. 20% of document\n2. 50% of document.\n\nSummary level: ?  ")
            if ratio=="1" or ratio=="20%" or ratio =="20":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.2))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize20.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                main()
            if ratio=="2" or ratio=="50%" or ratio =="50":
                print("\n")
                print ('Summary:')
                print (summarize(text_string, ratio=0.5))
                import docx
                doc = docx.Document()
                doc.add_paragraph(summarize(text_string, ratio=0.2))
                doc.save('Synopsize500.docx')
                toaster.show_toast("Synopsis saved.",
                                    " ",
                                    icon_path="cap.ico",
                                    duration=5)
                if __name__ == "__main__":
                    main()
            if __name__ == "__main__":
                main()
        if __name__ == "__main__":
            main()
    if __name__ == "__main__":
        main()
if __name__ == "__main__":
    main()
##        time.sleep(3)
##        print("\n")
##        input("Press Enter to exit Summarizer: ")
##        print("To synopsize document again, just start application again. ")
##        time.sleep(3)
##        sys.exit()
 
